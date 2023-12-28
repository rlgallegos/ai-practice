from docx import Document

# main_efforts = [
#     'Direct',
#     'Indirect',
#     'Strong',
#     'Light',
#     'Sudden',
#     'Sustained',
#     'Bound',
#     'Free'
# ]
effort_subs = [
    [
        'Direct',
        'Indirect'
    ],
    [
        'Strong',
        'Light'
    ],
    [
        'Sudden',
        'Sustained'
    ],
    [
        'Bound',
        'Free'
    ]
]

result = []
for space in effort_subs[0]:
    for weight in effort_subs[1]:
        for time in effort_subs[2]:
            for flow in effort_subs[3]:
                combo = space + ' - ' + weight + ' - ' + time + ' - ' + flow
                result.append(combo)

# efforts = [
#     'Direct Space',
#     'Indirect Space',
#     'Strong Weight',
#     'Light Weight',
#     'Sudden Time',
#     'Sustained Time',
#     'Bound Flow',
#     'Free Flow'
# ]

# actions = [
#     'Float',
#     'Punch',
#     'Glide',
#     'Slash',
#     'Dab',
#     'Wring',
#     'Flick',
#     'Press'
# ]



# for effort in efforts:
#     result.append('///')
#     for action in actions:
#         combo = effort + ' - ' + action
#         result.append(combo)

for r in result:
    print()
    print(r)

document = Document()

# Add a heading for the list
document.add_heading("Johnny's List", level=1)

# Create a paragraph and add the list items as bullet points
i = 0

for i in range(len(result)):
    if i % 8 == 0:
        paragraph = document.add_paragraph()
    paragraph.add_run(result[i] + "\n").font.name = "Times New Roman"
    

# Over each main effort
# for i in range(4):
#     effort = main_efforts[i]
#     document.add_heading(effort, level=1)

#     # Over each sub effort:
#     for j in range(2):
#         sub_effort = effort_subs[i][j]
#         document.add_heading(sub_effort, level=2)
#         paragraph = document.add_paragraph()

#         # Over each action
#         for k in range(8):
#             action = actions[k]
#             paragraph.add_run(action + "\u00a0\u00a0\u00a0").font.name = "Times New Roman"






# for item in result:
#     if i % 16 == 0 and i != 64:
#         print(i)
#         document.add_heading(main_efforts[i // 16])

#     if item == '///':
#         paragraph = document.add_paragraph()
#     elif i % 2 == 0:
#         paragraph.add_run(item + "\u00a0\u00a0\u00a0").font.name = "Times New Roman"
#     else:
#         paragraph.add_run(item + "\n").font.name = "Times New Roman"
#     i += 1

document.save("johnny-list.docx")